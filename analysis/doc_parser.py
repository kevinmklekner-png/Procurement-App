"""
Document parser for federal procurement solicitation documents.

Extracts text from DOCX/PDF/HTML, classifies document roles,
and parses structured SOW/PWS and evaluation criteria data.
"""

import io
import re
import json
from datetime import datetime
from typing import Dict, List, Optional


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file, including table content."""
    import docx
    doc = docx.Document(io.BytesIO(file_bytes))

    parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file using pdfplumber."""
    import pdfplumber
    parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                parts.append(text)
    return "\n".join(parts)


def extract_text_from_html(html: str) -> str:
    """Extract plain text from HTML."""
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(html, "html.parser")
    return soup.get_text(separator="\n", strip=True)


def extract_text(filename: str, content) -> str:
    """Auto-detect format and extract text.

    Args:
        filename: Original filename (used to detect format).
        content: bytes for PDF/DOCX, str for HTML.

    Returns:
        Extracted plain text.
    """
    lower = filename.lower()
    if lower.endswith(".docx"):
        return extract_text_from_docx(content)
    elif lower.endswith(".pdf"):
        return extract_text_from_pdf(content)
    elif lower.endswith((".html", ".htm")) or (isinstance(content, str) and "<" in content[:200]):
        html = content if isinstance(content, str) else content.decode("utf-8", errors="replace")
        return extract_text_from_html(html)
    # Fallback: try as plain text
    if isinstance(content, bytes):
        return content.decode("utf-8", errors="replace")
    return str(content)


# ---------------------------------------------------------------------------
# Document classification
# ---------------------------------------------------------------------------

SOW_SIGNALS = [
    "performance work statement", "statement of work", "scope of work",
    "labor categories", "labor category", "key personnel",
    "period of performance", "deliverables", "task order",
    "contract line item", "clin", "work requirements",
    "contractor shall", "the contractor shall",
]

EVAL_SIGNALS = [
    "evaluation criteria", "evaluation factor", "basis for award",
    "factor 1", "factor i", "evaluation methodology",
    "best value", "lowest price technically acceptable",
    "technical evaluation", "oral presentation",
    "past performance", "price evaluation",
    "adjectival rating", "rating scale",
]


def classify_document(filename: str, text_preview: str) -> str:
    """Classify a document based on filename and text content.

    Returns one of: sow, pws, solicitation, evaluation_criteria, unknown
    """
    lower_name = filename.lower()
    preview = text_preview[:5000].lower()

    # Filename-based hints
    if "pws" in lower_name:
        return "pws"
    if "sow" in lower_name and "rfp" not in lower_name:
        return "sow"

    # Score SOW signals
    sow_score = sum(1 for s in SOW_SIGNALS if s in preview)
    eval_score = sum(1 for s in EVAL_SIGNALS if s in preview)

    # Solicitation-specific signals (strong indicators of an RFP/solicitation doc)
    sol_signals = [
        "request for proposal", "rfp", "solicitation number",
        "section l", "section m", "offeror", "offerors",
        "phase i", "phase ii", "submission instructions",
        "representations and certifications", "evaluation and selection",
        "advisory down-select", "instructions to offerors",
    ]
    sol_score = sum(1 for s in sol_signals if s in preview)

    if sol_score >= 3:
        return "solicitation"
    if sol_score >= 1 and eval_score >= 2:
        return "solicitation"
    if eval_score >= 3 and eval_score > sow_score:
        return "evaluation_criteria"
    if sow_score >= 3 and sol_score == 0:
        return "sow"
    if sol_score >= 1:
        return "solicitation"

    return "unknown"


# ---------------------------------------------------------------------------
# SOW / PWS structured extraction
# ---------------------------------------------------------------------------

def _find_section(text: str, headers: List[str], stop_headers: Optional[List[str]] = None) -> str:
    """Find a section of text starting at any of *headers* and ending at the next major header."""
    for header in headers:
        pattern = rf"(?:^|\n)\s*(?:\d+[\.\)]\s*)?{re.escape(header)}\s*[:\.\-]?\s*\n"
        m = re.search(pattern, text, re.IGNORECASE)
        if m:
            start = m.end()
            # Find next section header (numbered or all-caps line)
            next_hdr = re.search(r"\n\s*(?:\d+[\.\)]\s+)?[A-Z][A-Z ]{5,}", text[start:])
            end = start + next_hdr.start() if next_hdr else start + 3000
            return text[start:end].strip()
    return ""


def _extract_labor_categories(text: str) -> List[str]:
    """Pull labor category names from text with noise filtering."""
    cats = []
    # Pattern: "Labor Category" header followed by bulleted/numbered list
    section = _find_section(text, ["labor categories", "labor category", "key personnel"])
    in_labeled_section = bool(section)
    if not section:
        section = text

    # Non-role words / section headers that look like role names but aren't
    noise = {
        "transition-out plan", "ordering procedures", "key personnel",
        "labor categories", "labor category", "deliverables",
        "background", "scope", "requirements", "overview",
        "submission instructions", "phase i", "phase ii",
        "far references", "volume", "attachment", "section",
        "contact", "list of", "evaluation", "instructions",
        "table of contents", "price quotation", "pricing",
        "invoice", "billing", "payment", "contract type",
        "task order", "base year", "option year", "period of performance",
        "place of performance", "government furnished", "contractor furnished",
        "quality assurance surveillance plan", "performance standards",
        "not applicable", "to be determined", "see below",
        "contractor support notification", "maintenance and support",
        "building operations support",
    }

    # Negative regex: reject lines that match these patterns
    _noise_re = re.compile(
        r"(?:"
        r"\b[A-Z]{2}\s+\d{5}"            # State + ZIP (e.g. "DC 20006")
        r"|FAR\s+\d+"                     # FAR clause references
        r"|DFARS\s+\d+"                   # DFARS clause references
        r"|^\d+\.\d+[\.\d]*\s"           # Section numbers like "1.2.3 "
        r"|^Row\s+\d+"                    # "Row N:" from forms
        r"|^\d{4}\s+[A-Z].*(?:Street|Ave|Blvd|Road|Suite|Floor)"  # Addresses
        r"|^(?:Phone|Fax|Email|Tel)[:\s]" # Contact info
        r"|^(?:Section|Part|Article)\s+[A-Z\d]"  # Section headers
        r"|^Attachment\s+[A-Z\d]"         # Attachment references
        r"|^(?:CLIN|SIN|FFP|T&M|LOE)\b"  # Contract line items
        r"|invoice|billing|payment|reimburs"  # Financial terms
        r"|BRANCH\s+\d"                       # Form fields like "IT SUPPORT BRANCH 01"
        r"|^Office of\b"                      # Institutional names
        r"|\bContracting Officer\b"           # CO references
        r"|Inspector General"                 # IG references
        r"|^(?:Inspect|Provide|Deliver|Install|Perform|Conduct|Ensure|Maintain|Review|Prepare)"  # Task verbs
        r"|workspace|office or|break room|conference room"  # Facility terms
        r"|name\s*plate|lightbulb|umbrella|sanitizer|refrigerator|microwave"  # Objects
        r"|^(?:Box|Paper|Picture)\s+\w+"  # Delivery/object lines
        r"|assistance$|delivery$|installation$|removal$|replacement$"  # Task noun endings
        r"|^\w+ing/\w+ing\b"             # "Opening/closing" pattern
        r"|^Scope of work\b|^Duration of\b|^Requestor\b|point of contact"  # SOW meta
        r")",
        re.IGNORECASE,
    )

    # Role words: at least one must appear (always required now)
    _role_words = re.compile(
        r"\b(?:manager|analyst|developer|engineer|specialist|consultant|"
        r"architect|lead|director|attorney|paralegal|associate|partner|"
        r"counsel|advisor|coordinator|designer|researcher|administrator|"
        r"expert|technician|scientist|auditor|accountant|writer|editor|"
        r"trainer|instructor|officer|planner|strategist|support|staff|"
        r"programmer|assistant|secretary|clerk|inspector|supervisor|"
        r"superintendent|foreman|mechanic|electrician|operator)\b",
        re.IGNORECASE,
    )

    # Look for lines that look like role names
    for line in section.split("\n"):
        line = line.strip().lstrip("•·-–—0123456789.) ")
        if not line or len(line) > 80:
            continue
        if line.lower() in noise:
            continue
        # Reject lines matching noise patterns
        if _noise_re.search(line):
            continue

        is_role_line = False

        # Typical labor cat patterns: "Senior Program Manager", "Subject Matter Expert (SME)"
        if re.match(r"^[A-Z][a-zA-Z /\-\(\)]+$", line) and 2 <= len(line.split()) <= 8:
            is_role_line = True
        # Also catch lines like "Program Manager – Level II"
        elif re.match(r"^[A-Z][a-zA-Z]+.*(?:Level|Sr|Jr|I{1,3}|Senior|Junior)", line) and len(line) <= 60:
            is_role_line = True

        if is_role_line:
            # Always require at least one role word
            if not _role_words.search(line):
                continue
            cats.append(line)

    # Also extract from inline mentions like "Senior Partner, Partner, Junior Associate"
    inline = re.search(
        r"(?:labor\s+categor\w+|key\s+personnel)\s+(?:shall\s+)?include[:\s]+(.*?)(?:\.|$)",
        text, re.IGNORECASE
    )
    if inline:
        for part in re.split(r",\s*(?:and\s+)?", inline.group(1)):
            part = part.strip()
            if part and 2 <= len(part.split()) <= 6 and len(part) <= 60:
                cats.append(part)

    return list(dict.fromkeys(cats))  # dedupe preserving order


def _extract_pop(text: str) -> str:
    """Extract period of performance info."""
    section = _find_section(text, ["period of performance"])
    if section:
        # Grab first couple of lines
        lines = [l.strip() for l in section.split("\n") if l.strip()][:3]
        return " ".join(lines)

    # Fallback: look for date range patterns
    m = re.search(
        r"(?:period of performance|pop)[:\s]*(.*?\d{4}.*?\d{4})",
        text, re.IGNORECASE
    )
    if m:
        return m.group(1).strip()
    return ""


def _extract_deliverables(text: str) -> List[str]:
    """Extract deliverable items."""
    section = _find_section(text, ["deliverables", "deliverable", "contract deliverables"])
    if not section:
        return []
    items = []
    for line in section.split("\n"):
        line = line.strip().lstrip("•·-–—0123456789.) ")
        if line and len(line) > 10:
            items.append(line)
    return items[:20]


def _extract_compliance(text: str) -> List[str]:
    """Extract compliance and regulatory requirements."""
    keywords = [
        "Section 508", "CUI", "Controlled Unclassified",
        "FISMA", "FedRAMP", "NIST 800", "ITAR", "EAR",
        "FAR 52", "DFARS", "Privacy Act", "HIPAA",
        "clearance", "security clearance", "public trust",
    ]
    found = []
    lower = text.lower()
    for kw in keywords:
        if kw.lower() in lower:
            found.append(kw)
    # Also pick up FAR clause references
    for m in re.finditer(r"FAR\s+\d+\.\d+[\-\d]*", text, re.IGNORECASE):
        found.append(m.group(0))
    return list(dict.fromkeys(found))


def _extract_key_tasks(text: str) -> List[str]:
    """Extract key tasks / task areas."""
    section = _find_section(text, [
        "scope of work", "scope", "tasks", "task areas",
        "requirements", "work requirements",
    ])
    if not section:
        return []
    tasks = []
    for line in section.split("\n"):
        line = line.strip().lstrip("•·-–—")
        # Numbered tasks like "1.1 Legal Research Support"
        m = re.match(r"^[\d\.]+\s+(.+)", line)
        if m:
            tasks.append(m.group(1).strip())
        elif line and len(line) > 15 and line[0].isupper():
            tasks.append(line)
    return tasks[:30]


def build_sow_analysis(document_id: int, solicitation_id: Optional[int],
                       notice_id: str, text: str) -> dict:
    """Build a structured SOW analysis dict matching the sow_analysis table."""
    scope_section = _find_section(text, [
        "scope", "scope of work", "introduction", "background and scope",
        "purpose", "objective",
    ])
    scope_summary = scope_section[:1000] if scope_section else ""

    labor_cats = _extract_labor_categories(text)
    pop = _extract_pop(text)
    deliverables = _extract_deliverables(text)
    compliance = _extract_compliance(text)
    key_tasks = _extract_key_tasks(text)

    # Place of performance
    pop_place = ""
    m = re.search(r"place of performance[:\s]*(.*?)(?:\n|$)", text, re.IGNORECASE)
    if m:
        pop_place = m.group(1).strip()

    # Confidence: higher if we found more fields
    fields_found = sum(bool(x) for x in [scope_summary, pop, labor_cats, deliverables, key_tasks])
    confidence = round(fields_found / 5.0, 2)

    return {
        "document_id": document_id,
        "solicitation_id": solicitation_id,
        "notice_id": notice_id,
        "scope_summary": scope_summary,
        "period_of_performance": pop,
        "place_of_performance": pop_place,
        "key_tasks": json.dumps(key_tasks),
        "labor_categories": json.dumps(labor_cats),
        "deliverables": json.dumps(deliverables),
        "compliance_reqs": json.dumps(compliance),
        "ordering_mechanism": "",
        "billing_instructions": "",
        "confidence_score": confidence,
        "extraction_method": "regex",
        "created_at": datetime.now().isoformat(),
    }


# ---------------------------------------------------------------------------
# Evaluation criteria extraction
# ---------------------------------------------------------------------------

def extract_evaluation_factors(document_id: int, solicitation_id: Optional[int],
                               notice_id: str, text: str) -> List[dict]:
    """Extract evaluation factors from solicitation text.

    Returns a list of dicts matching the evaluation_criteria table schema.
    """
    factors = []

    # Pattern 1: "Factor N: Name" or "Factor N - Name"
    factor_pattern = re.compile(
        r"(?:Factor|FACTOR)\s+(\d+|[IVX]+|[A-Z])\s*[:\-–—]\s*(.+?)(?:\n|$)",
        re.IGNORECASE
    )

    for m in factor_pattern.finditer(text):
        factor_num_raw = m.group(1).strip()
        factor_name = m.group(2).strip().rstrip(".")

        # Convert roman numerals or letters to number
        factor_number = _parse_factor_number(factor_num_raw)

        # Get surrounding text for description / subfactors
        start = m.end()
        next_factor = factor_pattern.search(text[start:])
        end = start + next_factor.start() if next_factor else start + 2000
        section = text[start:end]

        subfactors = _extract_subfactors(section)
        weight = _extract_weight(section, factor_name, text)
        page_limit = _extract_page_limit(section)
        rating_method = _extract_rating_method(section)
        phase = _extract_phase(text[:m.start()])

        factors.append({
            "document_id": document_id,
            "solicitation_id": solicitation_id,
            "notice_id": notice_id,
            "evaluation_phase": phase,
            "factor_number": factor_number,
            "factor_name": factor_name,
            "factor_weight": weight,
            "subfactors": json.dumps(subfactors),
            "description": section[:1000].strip(),
            "page_limit": page_limit,
            "rating_method": rating_method,
            "created_at": datetime.now().isoformat(),
        })

    # Pattern 2: If no "Factor N" found, try numbered list "1. Technical Approach"
    if not factors:
        numbered = re.compile(
            r"(?:^|\n)\s*(\d+)\.\s+([A-Z][A-Za-z /\-]+?)(?:\s*[\(\[]|[:\n])",
        )
        in_eval = False
        for line in text.split("\n"):
            if re.search(r"evaluation\s+(?:criteria|factor)", line, re.IGNORECASE):
                in_eval = True
            if in_eval:
                m = numbered.match(line.strip())
                if m:
                    factors.append({
                        "document_id": document_id,
                        "solicitation_id": solicitation_id,
                        "notice_id": notice_id,
                        "evaluation_phase": "",
                        "factor_number": int(m.group(1)),
                        "factor_name": m.group(2).strip(),
                        "factor_weight": "",
                        "subfactors": "[]",
                        "description": "",
                        "page_limit": "",
                        "rating_method": "",
                        "created_at": datetime.now().isoformat(),
                    })

    return factors


def _parse_factor_number(raw: str) -> int:
    """Convert a factor number string to int."""
    if raw.isdigit():
        return int(raw)
    # Roman numeral
    roman = {"I": 1, "II": 2, "III": 3, "IV": 4, "V": 5, "VI": 6, "VII": 7}
    if raw.upper() in roman:
        return roman[raw.upper()]
    # Letter
    if len(raw) == 1 and raw.isalpha():
        return ord(raw.upper()) - ord("A") + 1
    return 0


def _extract_subfactors(section: str) -> List[str]:
    """Extract subfactors from lettered lists like (a), (b), etc."""
    subs = []
    for m in re.finditer(r"\(([a-z])\)\s*(.+?)(?:\n|$)", section):
        subs.append(m.group(2).strip())
    if not subs:
        # Try "a. ", "b. " pattern
        for m in re.finditer(r"(?:^|\n)\s*([a-z])\.\s+(.+?)(?:\n|$)", section):
            subs.append(m.group(2).strip())
    return subs


def _extract_weight(section: str, factor_name: str, full_text: str) -> str:
    """Extract weight or relative importance description."""
    # Look for "most important", "equal", "descending order", percentages
    weight_patterns = [
        r"(\d+)\s*%",
        r"(most important|more important|equally important|descending order)",
        r"(significantly more important|slightly more important)",
        r"(greater|lesser|equal)\s+(?:weight|importance)",
    ]
    for pat in weight_patterns:
        m = re.search(pat, section, re.IGNORECASE)
        if m:
            return m.group(0).strip()

    # Check for ordering statements in the broader eval section
    m = re.search(
        r"(?:factors?|criteria)\s+(?:are|is)\s+(?:listed\s+)?in\s+descending\s+order\s+of\s+importance",
        full_text, re.IGNORECASE
    )
    if m:
        return "descending order of importance"
    return ""


def _extract_page_limit(section: str) -> str:
    """Extract page limit info."""
    m = re.search(r"(\d+)\s*(?:page|pg)s?\s*(?:limit|max|maximum)", section, re.IGNORECASE)
    if m:
        return f"{m.group(1)} pages"
    m = re.search(r"(?:limit|max|maximum)\s*(?:of\s*)?(\d+)\s*(?:page|pg)s?", section, re.IGNORECASE)
    if m:
        return f"{m.group(1)} pages"
    return ""


def _extract_rating_method(section: str) -> str:
    """Extract the rating / scoring method."""
    methods = [
        "adjectival", "color rating", "pass/fail", "pass fail",
        "acceptable/unacceptable", "go/no-go", "numerical score",
        "outstanding", "good", "acceptable", "marginal", "unacceptable",
    ]
    lower = section.lower()
    for method in methods:
        if method in lower:
            return method
    return ""


def _extract_phase(text_before: str) -> str:
    """Detect which evaluation phase we're in."""
    # Look backward for "Phase I" / "Phase II" markers
    phases = list(re.finditer(r"Phase\s+(I{1,3}|[123])", text_before, re.IGNORECASE))
    if phases:
        last = phases[-1].group(1)
        return f"Phase {last}"
    return ""


# ---------------------------------------------------------------------------
# CLI: parse on-disk files
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("Usage: python3 doc_parser.py <filename.docx|pdf|html>")
        sys.exit(1)

    filepath = sys.argv[1]
    with open(filepath, "rb") as f:
        content = f.read()

    filename = filepath.rsplit("/", 1)[-1]
    text = extract_text(filename, content)
    role = classify_document(filename, text)

    print(f"File: {filename}")
    print(f"Classified as: {role}")
    print(f"Text length: {len(text)} chars")
    print("=" * 60)

    if role in ("sow", "pws"):
        analysis = build_sow_analysis(0, None, "test", text)
        print("\n--- SOW/PWS ANALYSIS ---")
        print(f"Scope: {analysis['scope_summary'][:300]}...")
        print(f"PoP: {analysis['period_of_performance']}")
        print(f"Place: {analysis['place_of_performance']}")
        print(f"Labor Categories: {json.loads(analysis['labor_categories'])}")
        print(f"Key Tasks: {json.loads(analysis['key_tasks'])}")
        print(f"Deliverables: {json.loads(analysis['deliverables'])}")
        print(f"Compliance: {json.loads(analysis['compliance_reqs'])}")
        print(f"Confidence: {analysis['confidence_score']}")

    if role in ("solicitation", "evaluation_criteria"):
        factors = extract_evaluation_factors(0, None, "test", text)
        print(f"\n--- EVALUATION FACTORS ({len(factors)}) ---")
        for f in factors:
            print(f"  Factor {f['factor_number']}: {f['factor_name']}")
            print(f"    Weight: {f['factor_weight']}")
            subs = json.loads(f['subfactors'])
            if subs:
                print(f"    Subfactors: {subs}")
            if f['page_limit']:
                print(f"    Page limit: {f['page_limit']}")
            if f['rating_method']:
                print(f"    Rating: {f['rating_method']}")

    # If unknown or solicitation, try both parsers
    if role in ("unknown", "solicitation"):
        analysis = build_sow_analysis(0, None, "test", text)
        if json.loads(analysis['labor_categories']):
            print(f"\n--- Also found SOW content ---")
            print(f"Labor Categories: {json.loads(analysis['labor_categories'])}")
        factors = extract_evaluation_factors(0, None, "test", text)
        if factors:
            print(f"\n--- Also found {len(factors)} eval factors ---")
            for f in factors:
                print(f"  Factor {f['factor_number']}: {f['factor_name']}")
